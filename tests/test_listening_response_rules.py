from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from question_types import (  # noqa: E402
    ListeningResponseParser,
    parse_document_auto,
)
import wordtts as core  # noqa: E402


class ListeningResponseRuleTests(unittest.TestCase):
    @staticmethod
    def _make_document(path: Path) -> None:
        document = Document()
        paragraphs = [
            "七上 Starter Unit 1 Hello! 听后应答专项",
            "二、听后应答（共7小题）",
            "（计算机语音和屏幕文字提示）听句子，朗读正确应答语。",
            "（计算机语音提示）听下面1个句子。",
            "Good morning, class.",
            "（计算机语音提示）请朗读应答语。",
            "（计算机屏幕显示5秒倒计时进度条）",
            "★ Good morning, Peter.\t\t★ Good morning, Ms Gao.",
            "（计算机语音提示）听下面2个句子。",
            "First sentence.",
            "Second sentence.",
            "（计算机语音提示）请朗读应答语。",
            "（计算机语音提示）听下面1个句子。",
            "Inline one.\nInline two.",
            "（计算机语音提示）请朗读应答语。",
        ]
        for text in paragraphs:
            document.add_paragraph(text)
        document.save(path)

    def test_prompt_blocks_extract_each_content_line_and_ignore_answers(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "S1-听后应答.docx")
            self._make_document(path)
            result = ListeningResponseParser(path).parse()

        self.assertEqual(result["doc_type"], "听后应答")
        self.assertEqual(result["item_count"], 5)
        self.assertEqual(
            [item["filename_stem"] for item in result["items"]],
            [
                "7上-应答-1",
                "7上-应答-2",
                "7上-应答-3",
                "7上-应答-4",
                "7上-应答-5",
            ],
        )
        self.assertEqual(
            [item["text"] for item in result["items"]],
            [
                "Good morning, class.",
                "First sentence.",
                "Second sentence.",
                "Inline one.",
                "Inline two.",
            ],
        )
        self.assertTrue(all(item["voice"] == "female" for item in result["items"]))
        self.assertNotIn("Good morning, Peter", " ".join(item["text"] for item in result["items"]))

    def test_auto_detection_and_progress_keep_response_names(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "S1-听后应答.docx")
            self._make_document(path)
            results, summary = parse_document_auto(path)

        self.assertIn("检测到 1 种题型", summary)
        self.assertEqual([result["doc_type"] for result in results], ["听后应答"])
        progress = core.build_progress(
            path.name,
            str(path),
            results,
            {},
        )
        self.assertEqual(
            [item["filename"] for item in progress["items"]],
            [
                "7上-应答-1.mp3",
                "7上-应答-2.mp3",
                "7上-应答-3.mp3",
                "7上-应答-4.mp3",
                "7上-应答-5.mp3",
            ],
        )
        self.assertTrue(all(item["voice_override"] == "female" for item in progress["items"]))

    def test_multiple_word_paragraphs_are_not_truncated_by_prompt_count(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "九上-听后应答.docx")
            document = Document()
            for text in [
                "九上 听后应答",
                "（计算机语音提示）听下面1个句子。",
                "First response line.",
                "Second response line.",
                "（计算机语音提示）请朗读应答语。",
            ]:
                document.add_paragraph(text)
            document.save(path)
            result = ListeningResponseParser(path).parse()

        self.assertEqual(
            [item["filename_stem"] for item in result["items"]],
            ["9上-应答-1", "9上-应答-2"],
        )
        self.assertEqual(
            [item["text"] for item in result["items"]],
            ["First response line.", "Second response line."],
        )


if __name__ == "__main__":
    unittest.main()
