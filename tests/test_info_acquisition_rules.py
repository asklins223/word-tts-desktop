from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from question_types import InfoAcquisitionParser  # noqa: E402
import wordtts as core  # noqa: E402


class InfoAcquisitionQuestionRulesTests(unittest.TestCase):
    @staticmethod
    def _make_document(path: Path) -> None:
        document = Document()
        paragraphs = [
            "第一节 听选信息",
            "听第一段对话，回答第1—2两个问题。",
            "1. What is Amy doing?",
            "（Reading. / Singing. / Running.）",
            "2．Where is Tom?",
            "（At home. / At school. / In a park.）",
            "录音稿：",
            "W: Amy is reading.",
            "M: Tom is at school.",
            "第二节 回答问题",
            "3、Who helps Amy?",
            "4) Why is Tom happy?",
            "录音稿：",
            "(M) Their teacher helps them.",
            "参考答案：",
            "3. Their teacher.",
            "4. Because he finished his work.",
        ]
        for text in paragraphs:
            document.add_paragraph(text)
        document.save(path)

    def test_questions_continue_across_sections_and_alternate_voices(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "信息获取.docx")
            self._make_document(path)
            result = InfoAcquisitionParser(path).parse()

        questions = [
            item for item in result["items"] if item["category"].endswith("题目")
        ]
        self.assertEqual([item["number"] for item in questions], [1, 2, 3, 4])
        self.assertEqual(
            [item["category"] for item in questions],
            ["听选信息题目", "听选信息题目", "回答问题题目", "回答问题题目"],
        )
        self.assertEqual(
            [item["voice"] for item in questions],
            ["male", "female", "male", "female"],
        )
        self.assertEqual(
            [item["text"] for item in questions],
            [
                "M: What is Amy doing?",
                "W: Where is Tom?",
                "M: Who helps Amy?",
                "W: Why is Tom happy?",
            ],
        )
        self.assertEqual(
            [item["filename_stem"] for item in questions],
            ["问题1", "问题2", "问题3", "问题4"],
        )

        synthesized_segments = [core.parse_speakers(item["text"])[0] for item in questions]
        self.assertEqual(
            [voice for voice, _ in synthesized_segments],
            [
                core.MALE_VOICE,
                core.FEMALE_VOICE,
                core.MALE_VOICE,
                core.FEMALE_VOICE,
            ],
        )
        self.assertEqual(
            [text for _, text in synthesized_segments],
            [
                "What is Amy doing?",
                "Where is Tom?",
                "Who helps Amy?",
                "Why is Tom happy?",
            ],
        )

        scripts = [
            item for item in result["items"] if item["category"].endswith("录音稿")
        ]
        self.assertEqual(len(scripts), 2)

    @staticmethod
    def _make_document_new_format(path: Path) -> None:
        """新格式：每道题后紧跟行内「参考答案：xxx」，第二节问题漏编号，
        并带题号区间说明行「回答第 7-10 个问题」。"""
        document = Document()
        paragraphs = [
            "第一节 听选信息",
            "听第一段对话，回答第1—2两个问题。",
            "1. What is Amy doing?",
            "（Reading. / Singing. / Running.）",
            "参考答案：Reading./ She is reading.",
            "2．Where is Tom?",
            "（At home. / At school. / In a park.）",
            "参考答案：At home./ He is at home.",
            "录音稿：",
            "W: Amy is reading.",
            "M: Tom is at school.",
            "第二节 回答问题",
            "听下面一段独白，回答第 7-10 个问题。",
            "What is the speaker's favorite club?",
            "参考答案：The shadow play club.",
            "When do they practice?",
            "参考答案：Every Thursday.",
            "10. Where do they usually show?",
            "参考答案：On the stage.",
            "录音稿：",
            "(W) Script text.",
            "参考答案：",
            "7. The shadow play club.",
        ]
        for text in paragraphs:
            document.add_paragraph(text)
        document.save(path)

    def test_per_question_answers_and_missing_numbers(self):
        """行内「参考答案」不中断采集；漏编号题目按题号区间顺延。"""
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "信息获取.docx")
            self._make_document_new_format(path)
            result = InfoAcquisitionParser(path).parse()

        questions = [
            item for item in result["items"] if item["category"].endswith("题目")
        ]
        # 第一节 1-2；第二节说明行给出 7-10，漏编号题目顺延为 7、8，
        # 已编号 10 使用原题号
        self.assertEqual(
            [item["number"] for item in questions], [1, 2, 7, 8, 10]
        )
        self.assertEqual(
            [item["filename_stem"] for item in questions],
            ["问题1", "问题2", "问题7", "问题8", "问题10"],
        )
        self.assertEqual(
            [item["voice"] for item in questions],
            ["male", "female", "male", "female", "male"],
        )

        scripts = [
            item for item in result["items"] if item["category"].endswith("录音稿")
        ]
        self.assertEqual(len(scripts), 2)
        # 行内参考答案不得混入录音稿
        for script in scripts:
            self.assertNotIn("参考答案", script["text"])

    def test_recording_prompt_and_script_marker_split_adjacent_groups(self):
        """没有答案行时，下一组录音提示也必须切断上一段录音。"""
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "信息获取.docx")
            document = Document()
            for text in [
                "第一节 听选信息",
                "听下面一段录音，回答第1个问题。",
                "1. What is first?",
                "听力原文：",
                "W: First script.",
                "听下面一段录音，回答第2个问题。",
                "2. What is second?",
                "听力原文：",
                "M: Second script.",
            ]:
                document.add_paragraph(text)
            document.save(path)
            result = InfoAcquisitionParser(path).parse()

        scripts = [
            item for item in result["items"]
            if item["category"].endswith("录音稿")
        ]
        self.assertEqual(
            [item["text"] for item in scripts],
            ["W: First script.", "M: Second script."],
        )

    def test_script_marker_can_resume_after_standalone_answer_block(self):
        """答案区后紧跟「听力原文」时，不应静默丢掉后续录音。"""
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "信息获取.docx")
            document = Document()
            for text in [
                "第一节 听选信息",
                "1. What is it?",
                "参考答案：",
                "1. It is a book.",
                "听力原文：",
                "W: It is a book.",
            ]:
                document.add_paragraph(text)
            document.save(path)
            result = InfoAcquisitionParser(path).parse()

        scripts = [
            item for item in result["items"]
            if item["category"].endswith("录音稿")
        ]
        self.assertEqual([item["text"] for item in scripts], ["W: It is a book."])

    def test_inline_answer_after_script_does_not_start_a_fake_question(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "信息获取.docx")
            document = Document()
            for text in [
                "第一节 听选信息",
                "1. What is it?",
                "录音稿：",
                "W: It is a book.",
                "参考答案：It is a book.",
                "1. It is a book.",
                "听下面一段录音，回答第2个问题。",
                "2. What is next?",
                "录音稿：",
                "M: The next answer.",
            ]:
                document.add_paragraph(text)
            document.save(path)
            result = InfoAcquisitionParser(path).parse()

        questions = [
            item for item in result["items"]
            if item["category"].endswith("题目")
        ]
        scripts = [
            item for item in result["items"]
            if item["category"].endswith("录音稿")
        ]
        self.assertEqual([item["number"] for item in questions], [1, 2])
        self.assertEqual(
            [item["text"] for item in scripts],
            ["W: It is a book.", "M: The next answer."],
        )

    def test_question_audio_files_use_question_numbers(self):
        parse_results = [{
            "doc_type": "信息获取",
            "items": [
                {
                    "category": "听选信息题目",
                    "number": 1,
                    "filename_stem": "问题1",
                    "text": "M: First question?",
                },
                {
                    "category": "回答问题题目",
                    "number": 7,
                    "filename_stem": "问题7",
                    "text": "M: Seventh question?",
                },
                {
                    "category": "回答问题录音稿",
                    "index": 1,
                    "text": "(W) Script text.",
                },
            ],
        }]

        progress = core.build_progress(
            "信息获取.docx",
            "/tmp/信息获取.docx",
            parse_results,
            {"format": "mp3"},
        )

        self.assertEqual(
            [item["filename"] for item in progress["items"]],
            ["问题1.mp3", "问题7.mp3", "回答问题-录音稿1.mp3"],
        )


if __name__ == "__main__":
    unittest.main()
