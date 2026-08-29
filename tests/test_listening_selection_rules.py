from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from question_types import (  # noqa: E402
    ListeningSelectionParser,
    parse_document_auto,
)
import wordtts as core  # noqa: E402


class ListeningSelectionRuleTests(unittest.TestCase):
    @staticmethod
    def _make_document(path: Path) -> None:
        document = Document()
        paragraphs = [
            "7上 Starter Unit 1 Hello!",
            "一、听后选择（共4小题）",
            "（计算机语音提示）听下面一段对话，回答第1小题。",
            "1. How is Peter?",
            "A. He is tired.",
            "B. He is fine.",
            "【录音原文】",
            "W: Good morning, Peter! How are you?",
            "M: I’m fine, thanks.",
            "（计算机语音提示）听下面一段对话，回答第2小题。",
            "2. What is the girl's name?",
            "【录音原文】",
            "M: Hello! I’m Jack.",
            "W: My name is Emma.",
            "二、其他题型",
            "【录音原文】",
            "W: This must not be collected.",
        ]
        for text in paragraphs:
            document.add_paragraph(text)
        document.save(path)

    @staticmethod
    def _make_mixed_document(path: Path) -> None:
        document = Document()
        paragraphs = [
            "第一节 听选信息",
            "1. What is the answer?",
            "录音稿：",
            "W: Old question script.",
            "一、听后选择（共1小题）",
            "（计算机语音提示）听下面一段对话。",
            "1. What is new?",
            "A. This.",
            "B. That.",
            "C. Other.",
            "【录音原文】",
            "m: New question script.",
            "二、其他题型",
            "录音稿：",
            "W: Must not leak into either result.",
        ]
        for text in paragraphs:
            document.add_paragraph(text)
        document.save(path)

    def test_only_recording_dialogues_are_extracted(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "听后选择.docx")
            self._make_document(path)
            result = ListeningSelectionParser(path).parse()

        self.assertEqual(result["doc_type"], "听后选择")
        self.assertEqual(result["item_count"], 2)
        self.assertEqual(
            [item["category"] for item in result["items"]],
            ["听后选择录音稿"] * 2,
        )
        self.assertEqual(
            [item["text"] for item in result["items"]],
            [
                "W: Good morning, Peter! How are you?\nM: I’m fine, thanks.",
                "M: Hello! I’m Jack.\nW: My name is Emma.",
            ],
        )
        self.assertEqual(
            [item["question_index"] for item in result["items"]],
            [1, 2],
        )
        self.assertNotIn("How is Peter", result["items"][0]["text"])
        self.assertNotIn("计算机语音提示", result["items"][0]["text"])

    def test_content_detection_and_wm_labels_are_supported(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "7上 Starter Unit 1 Hello!.docx")
            self._make_document(path)
            results, summary = parse_document_auto(path)

        self.assertIn("检测到 1 种题型", summary)
        self.assertEqual([result["doc_type"] for result in results], ["听后选择"])
        self.assertEqual(results[0]["item_count"], 2)

        first_segments = core.build_synthesis_segments(
            results[0]["items"][0]["text"],
            50,
            50,
            50,
        )
        self.assertEqual(
            [(segment["voice_key"], segment["text"]) for segment in first_segments],
            [
                (core.FEMALE_VOICE, "Good morning, Peter! How are you?"),
                (core.MALE_VOICE, "I’m fine, thanks."),
            ],
        )

        second_segments = core.build_synthesis_segments(
            results[0]["items"][1]["text"],
            50,
            50,
            50,
        )
        self.assertEqual(
            [(segment["voice_key"], segment["text"]) for segment in second_segments],
            [
                (core.MALE_VOICE, "Hello! I’m Jack."),
                (core.FEMALE_VOICE, "My name is Emma."),
            ],
        )

        lowercase_segments = core.build_synthesis_segments(
            "w: Female line.\nm: Male line.",
            50,
            50,
            50,
        )
        self.assertEqual(
            [(segment["voice_key"], segment["text"]) for segment in lowercase_segments],
            [
                (core.FEMALE_VOICE, "Female line."),
                (core.MALE_VOICE, "Male line."),
            ],
        )

    def test_each_recording_block_stays_one_progress_item(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "听后选择.docx")
            self._make_document(path)
            results = [ListeningSelectionParser(path).parse()]

        progress = core.build_progress(
            "听后选择.docx",
            str(path),
            results,
            {"generation_mode": core.GENERATION_MODE_COMPOSITE},
        )

        self.assertEqual(progress["total_items"], 2)
        self.assertEqual(
            [item["filename"] for item in progress["items"]],
            ["听后选择-录音稿1.mp3", "听后选择-录音稿2.mp3"],
        )
        self.assertEqual(
            [item["raw_item"]["text"] for item in progress["items"]],
            [
                "W: Good morning, Peter! How are you?\nM: I’m fine, thanks.",
                "M: Hello! I’m Jack.\nW: My name is Emma.",
            ],
        )

    def test_new_type_does_not_leak_into_old_type_in_a_mixed_document(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "新旧题型混合.docx")
            self._make_mixed_document(path)
            results, summary = parse_document_auto(path)

        self.assertIn("检测到 2 种题型", summary)
        by_type = {result["doc_type"]: result for result in results}
        self.assertEqual(
            [item["text"] for item in by_type["信息获取"]["items"]
             if item["category"].endswith("录音稿")],
            ["W: Old question script."],
        )
        self.assertEqual(
            [item["text"] for item in by_type["听后选择"]["items"]],
            ["m: New question script."],
        )


if __name__ == "__main__":
    unittest.main()
