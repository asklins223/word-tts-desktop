from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "word_parser"))

from word_parser import TextReadingParser  # noqa: E402


class TextReadingRuleTests(unittest.TestCase):
    @staticmethod
    def _save_new_format(
        path: Path,
        title: str = "课文跟读新格式",
        include_conversations: bool = True,
    ) -> None:
        document = Document()
        dialogue = [
            "Conversation 2",
            "Bob: Bob's line.",
            "Alice: Alice's line.",
            "Conversation 1",
            "Teacher: Hello, class.",
            "Class: Hello, teacher.",
        ]
        if not include_conversations:
            dialogue = [
                "Bob: Bob's line.",
                "Alice: Alice's line.",
                "Teacher: Hello, class.",
                "Class: Hello, teacher.",
            ]
        paragraphs = [
            title,
            "Section A",
            "句子跟读",
            "2. Second sentence.",
            "中文：第二句。",
            "1. First sentence.",
            "中文：第一句。",
            "段落跟读",
            *dialogue,
            "Section B",
            "语篇跟读：",
            "// Welcome",
            "A short introduction.",
            "The first story",
            "The first story text is here.",
            "A second paragraph follows.",
        ]
        for text in paragraphs:
            document.add_paragraph(text)
        document.save(path)

    @staticmethod
    def _save_legacy_format(path: Path) -> None:
        document = Document()
        document.add_heading("Understanding Idea", level=1)
        document.add_heading("句子跟读", level=2)
        document.add_paragraph("1. One sentence.")
        document.add_heading("段落跟读", level=2)
        document.add_paragraph("First sentence. Second sentence.")
        document.add_heading("语篇跟读", level=2)
        document.add_paragraph("语篇1")
        document.add_paragraph("First discourse sentence. Second discourse sentence.")
        document.save(path)

    @staticmethod
    def _save_role_discourse_document(path: Path) -> None:
        document = Document()
        for text in [
            "课程跟读-角色结构样本",
            "Section B",
            "语篇跟读",
            "Conversation 1",
            "Teng Fei: Good morning.",
            "Emma: Good morning, Teng Fei.",
        ]:
            document.add_paragraph(text)
        document.save(path)

    def test_section_format_uses_structure_for_dialogue_mode_and_article_boundaries(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "课文跟读新格式.docx")
            self._save_new_format(path, include_conversations=False)
            parser = TextReadingParser(path)
            result = parser.parse()

        sentences = [item for item in result["items"] if item["category"] == "句子跟读"]
        paragraphs = [item for item in result["items"] if item["category"] == "段落跟读"]
        discourses = [item for item in result["items"] if item["category"] == "语篇跟读"]

        self.assertEqual([item["filename_stem"] for item in sentences], ["SA句子1", "SA句子2"])
        self.assertEqual([item["text"] for item in sentences], ["First sentence.", "Second sentence."])
        self.assertTrue(all(item["voice"] == "female" for item in sentences))

        self.assertEqual([item["filename_stem"] for item in paragraphs], ["SA段落1"])
        self.assertEqual(parser._detect_section_ab_profile()["role_audio_mode"], "aggregate")
        self.assertEqual(
            paragraphs[0]["text"],
            "Bob: Bob's line.\nAlice: Alice's line.\n"
            "Teacher: Hello, class.\nClass: Hello, teacher.",
        )

        self.assertEqual(
            [item["filename_stem"] for item in discourses],
            ["SB语篇1", "SB语篇2", "SB语篇3"],
        )
        self.assertEqual(discourses[0]["text"], "Welcome\nA short introduction.")
        self.assertEqual(
            discourses[1]["text"],
            "The first story\nThe first story text is here.",
        )
        self.assertEqual(discourses[2]["text"], "A second paragraph follows.")

    def test_legacy_format_keeps_legacy_sentence_splitting(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "课文跟读旧格式.docx")
            self._save_legacy_format(path)
            result = TextReadingParser(path).parse()

        paragraphs = [item for item in result["items"] if item["category"] == "段落跟读"]
        discourses = [item for item in result["items"] if item["category"] == "语篇跟读"]

        self.assertEqual(len(paragraphs), 2)
        self.assertEqual([item["filename_stem"] for item in paragraphs], ["U-段落1", "U-段落2"])
        self.assertEqual(len(discourses), 2)
        self.assertEqual(
            [item["filename_stem"] for item in discourses],
            ["U-语篇1-1", "U-语篇1-2"],
        )

    def test_conversation_structure_splits_each_role_into_an_audio_and_exposes_role(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "课文跟读-对话结构样本.docx")
            self._save_new_format(path, title="课程对话结构样本")
            parser = TextReadingParser(path)
            result = parser.parse()

        paragraphs = [item for item in result["items"] if item["category"] == "段落跟读"]
        self.assertEqual(parser._detect_section_ab_profile()["role_audio_mode"], "per_role")
        self.assertEqual(
            [item["filename_stem"] for item in paragraphs],
            ["SA段落1", "SA段落2", "SA段落3", "SA段落4"],
        )
        self.assertEqual(
            [item["role"] for item in paragraphs],
            ["Bob", "Alice", "Teacher", "Class"],
        )
        self.assertEqual(
            [item["text"] for item in paragraphs],
            [
                "Bob: Bob's line.",
                "Alice: Alice's line.",
                "Teacher: Hello, class.",
                "Class: Hello, teacher.",
            ],
        )

    def test_conversation_discourse_also_splits_each_role(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "课文跟读-角色对话.docx")
            self._save_role_discourse_document(path)
            parser = TextReadingParser(path)
            result = parser.parse()

        discourses = [item for item in result["items"] if item["category"] == "语篇跟读"]
        self.assertEqual(parser._detect_section_ab_profile()["role_audio_mode"], "per_role")
        self.assertEqual([item["filename_stem"] for item in discourses], ["SB语篇1", "SB语篇2"])
        self.assertEqual([item["role"] for item in discourses], ["Teng Fei", "Emma"])

    def test_isolated_section_marker_does_not_switch_to_new_parser(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "课文跟读旧格式含备注.docx")
            document = Document()
            for text in [
                "Section A",
                "Section A is mentioned in the teacher's note.",
                "Understanding Idea",
                "句子跟读",
                "1. Legacy sentence.",
            ]:
                document.add_paragraph(text)
            document.save(path)

            parser = TextReadingParser(path)
            self.assertFalse(parser._is_section_ab_format())
            result = parser.parse()

        self.assertEqual(result["items"][0]["filename_stem"], "U-句子1")

    def test_new_structure_wins_when_old_and_new_sections_are_both_present(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "课文跟读混合结构.docx")
            document = Document()
            for text in [
                "Understanding Idea",
                "句子跟读",
                "1. Old sentence.",
                "Section A",
                "句子跟读",
                "1. New sentence.",
            ]:
                document.add_paragraph(text)
            document.save(path)

            parser = TextReadingParser(path)
            self.assertTrue(parser._is_section_ab_format())
            result = parser.parse()

        self.assertEqual(
            [item["filename_stem"] for item in result["items"]],
            ["SA句子1"],
        )
        self.assertEqual(result["items"][0]["text"], "New sentence.")

    def test_reading_plus_uses_rp_prefix_and_inline_slash_marker(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir, "课文跟读人教九上U1.docx")
            document = Document()
            for text in [
                "课文跟读人教九上U1",
                "Section B",
                "Reading Plus",
                "语篇跟读",
                "RP title",
                "//Small title\nSmall title paragraph.",
                "Last paragraph.",
            ]:
                document.add_paragraph(text)
            document.save(path)
            result = TextReadingParser(path).parse()

        discourses = [item for item in result["items"] if item["category"] == "语篇跟读"]
        self.assertEqual(
            [item["filename_stem"] for item in discourses],
            ["RP语篇1", "RP语篇2", "RP语篇3"],
        )
        self.assertTrue(all(item["section"] == "Reading Plus" for item in discourses))
        self.assertEqual(discourses[1]["text"], "Small title\nSmall title paragraph.")


if __name__ == "__main__":
    unittest.main()
